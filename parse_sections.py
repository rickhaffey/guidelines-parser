import docx
import re

sectionRegexes = {
  r"""secondary.*assessment.*treatment.*interventions""": "Secondary Assessment, Treatment, and Interventions",
  r"""assessment.*treatment.*interventions""": "Assessment, Treatment, and Interventions",
  r"""treatment.*interventions""": "Treatment and Interventions",
  r"""assessment""": "Assessment",
  r"""definitions""": "Definitions",
  r"""inclusion.*exclusion.*criteria""": "Inclusion / Exclusion Criteria",
  r"""exclusion.*criteria""": "Exclusion Criteria",
  r"""inclusion.*criteria""": "Inclusion Criteria",
  r"""key.*considerations""": "Key Considerations",
  r"""key.*documentation.*elements""": "Key Documentation Elements",
  r"""notes.*educational.*pearls""": "Notes / Educational Pearls",
  r"""patient.*care.*goals""": "Patient Care Goals",
  r"""patient.*management""": "Patient Management",
  r"""patient.*presentation""": "Patient Presentation",
  r"""patient.*safety.*considerations""": "Patient Safety Considerations",
  r"""performance.*measures""": "Performance Measures",
  r"""pertinent.*assessment.*findings""": "Pertinent Assessment Findings",
  r"""quality.*improvement""": "Quality Improvement",
  r"""references""": "References",
  r"""special.*transport.*considerations""": "Special Transport Considerations",
  r"""scene.*management""": "Scene Management"
}


def get_section_header(p):
    rtext = ""
    isSection = False
    for r in p.runs:
        rtext = rtext + r.text.strip()
        f = r.font
        isSection = isSection or (f.bold and f.underline)

    if(isSection):
        # check to see whether it matches one of our regexes
        for regex in sectionRegexes.keys():
            if(re.match(regex, rtext, re.I)):
                return sectionRegexes[regex]

    return None


def main():
    doc = docx.Document('./guidelines.docx')

    for i in range(0, len(doc.paragraphs)):
        p = doc.paragraphs[i]

        if(p.text == "APPENDICES"):
            break

        header = get_section_header(p)
        if(header):
            print(header)

    
if __name__ == '__main__':
    main()
