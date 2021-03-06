import docx
import json
import re
import sys

default_root = "/Users/rhaffey/Dropbox/Projects/EMS/EBG"

json_category = "category"
json_guidelines = "guidelines"
json_guideline = "guideline"
json_title = "title"
json_sections = "sections"
json_p_index = "p_index"
json_heading = "heading"
json_text = "text"
json_items = "items"
json_indent = "indent"
json_nemsis = "nemsis_refs"
json_nemsis_id = "id"
json_nemsis_text = "text"

# json section names
json_s_defintions = "defintions"
json_s_patientCareGoals = "patientCareGoals"
json_s_patientPresentation = "patientPresentation"
json_s_inclusionCriteria = "inclusionCriteria"
json_s_exclusionCriteria = "exclusionCriteria"
json_s_inclusionExclusionCriteria = "inclusionExclusionCriteria"
json_s_specialTransporConsiderations = "specialTransporConsiderations"
json_s_sceneManagement = "sceneManagement"
json_s_patientManagement = "patientManagement"
json_s_assessment = "assessment"
json_s_treatmentAndInterventions = "treatmentAndInterventions"
json_s_patientSafetyConsiderations = "patientSafetyConsiderations"
json_s_secondaryAssesmentTreatmentAndInterventions = "secondaryAssesmentTreatmentAndInterventions"
json_s_assessmentTreatmentAndInterventions = "assessmentTreatmentAndInterventions"
json_s_notesAndEducationalPearls = "notesAndEducationalPearls"
json_s_keyConsiderations = "keyConsiderations"
json_s_pertinentAssessmentFindings = "pertinentAssessmentFindings"
json_s_qualityImprovement = "qualityImprovement"
json_s_keyDocumentationElements = "keyDocumentationElements"
json_s_performanceMeasures = "performanceMeasures"
json_s_references = "references"
json_s_revisionDate = "revisionDate"


sectionRegexes = [
  (r"""secondary.*assessment.*treatment.*interventions""", [json_s_patientManagement, json_s_secondaryAssesmentTreatmentAndInterventions]),
  (r"""assessment.*treatment.*interventions""", [json_s_patientManagement, json_s_assessmentTreatmentAndInterventions]),
  (r"""treatment.*interventions""", [json_s_patientManagement, json_s_treatmentAndInterventions]),
  (r"""assessment""", [json_s_patientManagement, json_s_assessment]),
  (r"""definitions""", [json_s_defintions]),
  (r"""inclusion.*exclusion.*criteria""", [json_s_patientPresentation, json_s_inclusionExclusionCriteria]),
  (r"""exclusion.*criteria""", [json_s_patientPresentation, json_s_exclusionCriteria]),
  (r"""inclusion.*criteria""", [json_s_patientPresentation, json_s_inclusionCriteria]),
  (r"""key.*considerations""", [json_s_notesAndEducationalPearls, json_s_keyConsiderations]),
  (r"""key.*documentation.*elements""", [json_s_qualityImprovement, json_s_keyDocumentationElements]),
  (r"""notes.*educational.*pearls""", [json_s_notesAndEducationalPearls]),
  (r"""patient.*care.*goals""", [json_s_patientCareGoals]),
  (r"""patient.*management""", [json_s_patientManagement]),
  (r"""patient.*presentation""", [json_s_patientPresentation]),
  (r"""patient.*safety.*considerations""", [json_s_patientManagement, json_s_patientSafetyConsiderations]),
  (r"""performance.*measures""", [json_s_qualityImprovement, json_s_performanceMeasures]),
  (r"""pertinent.*assessment.*findings""", [json_s_notesAndEducationalPearls, json_s_pertinentAssessmentFindings]),
  (r"""quality.*improvement""", [json_s_qualityImprovement]),
  (r"""references""", [json_s_references]),
  (r"""special.*transport.*considerations""", [json_s_specialTransporConsiderations]),
  (r"""scene.*management""", [json_s_sceneManagement]),
  (r"""revision.*date""", [json_s_revisionDate])
]


def build_category(text):
    return (json_category, text.strip())


def build_guideline(text, i, category):
    return (json_guideline, {
        json_title: text.strip(),
        json_category: category,
        # sections - will be added dynamically as they're found in the source doc
        json_nemsis: None,
        json_p_index: i
    })

def build_nemsis_ref(text, i):
    # parse out the parts we care about
    # todo - we're using the regex 2x (in the is_nemsis method and here) - find a clean way to do this once?
    # note - restrict to cases with 3 or more digits following leading paren.
    regex = r"""^\((\d{3,}\W*.*)\)"""
    match = re.match(regex, text, re.I)

    if(match != None and len(match.groups()) >= 1):
        # have a match; split on ';', then parse ID and value for each
        results = []
        reftext = match.groups()[0]
        refs = [r.split('–') for r in reftext.split(';')]

        for ref in refs:
            refid = int(ref[0].strip())
            reftext = ref[1].strip()
            results.append({
                json_nemsis_id: refid,
                json_nemsis_text: reftext,
                json_p_index: i
                })

        return (json_nemsis, results)
    else:
        return (json_nemsis, None)


def build_section(i):
    return {
        json_p_index: i,
        json_items: []
    }


def build_section_text(text, i):
    return {
        json_text: text,
        json_p_index: i
    }


def is_heading1(style):
    return ("Heading1" in style or "Heading 1" in style)


def is_heading2(style):
    return ("Heading2" in style or "Heading 2" in style)


def is_nemsis_ref(text):
    regex = r"""^\(\d+"""
    return re.match(regex, text, re.I)


def parse_paragraph(p, i, category = None):
    style = p.style.name

    if(is_heading1(style)):
        return build_category(p.text)
    elif(is_heading2(style)):
        return build_guideline(p.text, i, category)
    elif(is_nemsis_ref(p.text)):
        return build_nemsis_ref(p.text, i)

    runheading = ""
    for r in p.runs:
        style = r.style.name
        if("Heading" in style):
            runheading = runheading + r.text

    if(runheading.strip()):
        return build_guideline(runheading, i, category)

    return None


def parse_section_header(p):
    rtext = ""
    isSection = False
    for r in p.runs:
        rtext = rtext + r.text.strip()
        f = r.font
        isSection = isSection or (f.bold and f.underline)

    if(isSection):
        # check to see whether it matches one of our regexes
        for (regex, hierarchy) in sectionRegexes:
            if(re.match(regex, rtext, re.I)):
                return hierarchy

    return None


def get_ppr(p):
  """Gets the product properties (serialized as <pPr> in the .docx xml) for the paragraph"""
  return p._element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")


def get_ilvl(p):
  """Gets the indent level (serialized as <pPr><numPr><ilvl val="_"> in the .docx xml) from the product properties"""
  ppr = get_ppr(p)

  if(ppr is None or ppr.numPr is None or ppr.numPr.ilvl is None or ppr.numPr.ilvl.val is None):
    return 0
  return int(ppr.numPr.ilvl.val) + 1


def show_progress(marker):
    print(marker, end='', flush=True)


abridged = False

def get_infile_path():
    if(len(sys.argv) >= 2):
        return sys.argv[1]
    else:
        if(abridged):
            return default_root + "/guidelines.partial.docx"
        else:
            return default_root + "/guidelines.docx"


def get_outfile_path():
    if(len(sys.argv) >= 3):
        return sys.argv[2]
    else:
        if(abridged):
            return default_root + "/output/out.partial.json"
        else:
            return default_root + "/output/out.json"


def main():
    # read in the input word doc
    path = get_infile_path()
    source_doc = docx.Document(path)

    # build the starting point of an output doc
    out_doc = {
        json_guidelines: []
    }

    # build some preliminary 'running' objects
    running_category = None
    running_guideline = None
    running_section = None

    # iterate over the paragraphs in the doc
    start = 0
    end = start + len(source_doc.paragraphs)
    for i in range(start, end):
        p = source_doc.paragraphs[i]

        # debugging
        # print("%d: %s" % (i, p.text))

        # first, handle new category and guideline paragraphs
        parsed = parse_paragraph(p, i, running_category)
        if(parsed):
            if(parsed[0] == json_category):
                running_category = parsed[1]
                show_progress("\n# " + running_category)
            elif(parsed[0] == json_guideline):
                running_guideline = parsed[1]
                out_doc[json_guidelines].append(running_guideline)
                show_progress("\n\t+")
            elif(parsed[0] == json_nemsis):
                running_guideline[json_nemsis] = parsed[1]
                show_progress("n")

        # next, parse out any section details (headers, and section text)
        section_path = parse_section_header(p)
        if(section_path):
            # iterate through the elements of the path, creating any if they're not found
            parent = running_guideline
            for section_path_element in section_path:
                element = parent.get(section_path_element)
                if(element is None):
                    element = build_section(i)
                    parent[section_path_element] = element

                parent = element

            running_section = element
            item_parents = {}
        elif(parsed == None and running_section != None):
            outText = p.text.strip()
            if(outText != ""):
                lvl = get_ilvl(p)
                section_text = build_section_text(outText, i)
                item_parents[lvl] = section_text
                item_parent = item_parents.get(lvl - 1)
                if(item_parent == None):
                    running_section[json_items].append(section_text)
                else:
                    if(item_parent.get(json_items) == None):
                        item_parent[json_items] = []
                    item_parent[json_items].append(section_text)
                show_progress(".")
            else:
                show_progress("x")

    # write the final document in .json format to the outfile
    with open(get_outfile_path(), 'w') as outfile:
        outfile.write(json.dumps(out_doc, indent=2))


if __name__ == '__main__':
    main()
